"""
Dynamic Remediation Preview engine.

Given the *original* uploaded document (its real bytes on disk) and the list
of violations that were just resolved, this module produces a new version of
that same document with only the offending excerpts swapped for their
AI-compliant rewrites. It never uses a template or placeholder file — it
edits a copy of the actual uploaded document, so formatting, layout, images,
tables, and all untouched content are preserved.

Supported formats:
  - .docx : edited in place with python-docx (paragraphs, tables, headers,
            footers). Only the runs containing a flagged excerpt are
            rewritten; every other run/paragraph/image/table is untouched.
  - .pdf  : edited in place with PyMuPDF (fitz) using text redaction +
            re-insertion at the same location, which preserves the original
            page layout/images far better than regenerating the PDF from
            scratch.
  - .txt / .md : plain string replacement.
"""
from __future__ import annotations

import io
from pathlib import Path
from typing import List, Tuple
from difflib import SequenceMatcher
import re

def find_source_span(text: str, excerpt: str) -> str | None:
    """
    Given a (possibly LLM-paraphrased) excerpt, return the actual verbatim
    substring of `text` it refers to, using the same match strategies as
    smart_replace (exact -> whitespace-normalized -> sentence fuzzy -> word-
    window fuzzy). Returns None if no reasonable match is found.

    This exists so that a Violation's `excerpt` can be "snapped" back to real
    document text at detection time (see agents.py), guaranteeing that the
    later document-regeneration replacement step is always operating on text
    that genuinely exists in the document — never silently failing to find
    a match because the LLM slightly reworded the excerpt it reported.
    """
    if not excerpt or not text:
        return None

    if excerpt in text:
        return excerpt

    def normalize(s: str) -> str:
        return re.sub(r"\s+", " ", s.strip())

    norm_text = normalize(text)
    norm_excerpt = normalize(excerpt)

    if norm_excerpt in norm_text:
        idx = norm_text.index(norm_excerpt)
        orig_idx = 0
        norm_idx = 0
        while norm_idx < idx and orig_idx < len(text):
            if text[orig_idx] in " \n\t\r":
                while orig_idx < len(text) and text[orig_idx] in " \n\t\r":
                    orig_idx += 1
                norm_idx += 1
            else:
                orig_idx += 1
                norm_idx += 1
        end_idx = orig_idx
        covered = 0
        while covered < len(norm_excerpt) and end_idx < len(text):
            if text[end_idx] in " \n\t\r":
                while end_idx < len(text) and text[end_idx] in " \n\t\r":
                    end_idx += 1
                covered += 1
            else:
                end_idx += 1
                covered += 1
        span = text[orig_idx:end_idx]
        if span:
            return span

    sentences = re.split(r"(?<=[.!?\n])\s*", text)
    best_match, best_ratio = None, 0.0
    for sentence in sentences:
        if not sentence.strip():
            continue
        ratio = SequenceMatcher(None, normalize(sentence), norm_excerpt).ratio()
        if ratio > best_ratio:
            best_ratio, best_match = ratio, sentence
    if best_match and best_ratio >= 0.75:
        return best_match

    window_size = max(1, len(norm_excerpt.split()))
    words = text.split()
    best_ratio, best_start = 0.0, -1
    for i in range(len(words) - window_size + 1):
        window = " ".join(words[i : i + window_size])
        ratio = SequenceMatcher(None, normalize(window), norm_excerpt).ratio()
        if ratio > best_ratio:
            best_ratio, best_start = ratio, i
    if best_ratio >= 0.70 and best_start >= 0:
        return " ".join(words[best_start : best_start + window_size])

    return None


def smart_replace(text: str, excerpt: str, replacement: str) -> str:
    """
    Replace a clause even if spacing, newlines or punctuation differ slightly.
    Tries four strategies in order of precision:
    1. Exact match
    2. Normalised-whitespace exact match
    3. Sentence-level fuzzy match (SequenceMatcher ≥ 0.85)
    4. Word-token fuzzy match (SequenceMatcher ≥ 0.80) as last resort
    """

    if not excerpt or not text:
        return text

    # 1. Exact match (fast path)
    if excerpt in text:
        return text.replace(excerpt, replacement, 1)

    def normalize(s: str) -> str:
        """Collapse all whitespace (including newlines) to single spaces."""
        return re.sub(r"\s+", " ", s.strip())

    norm_text    = normalize(text)
    norm_excerpt = normalize(excerpt)

    # 2. Normalised-whitespace exact match
    if norm_excerpt in norm_text:
        # Find the original span in `text` that corresponds to norm_excerpt
        # by locating the first character and scanning forward.
        idx = norm_text.index(norm_excerpt)
        # Map the normalised index back to a rough char position in the
        # original text using a simple scan.
        orig_idx = 0
        norm_idx = 0
        while norm_idx < idx and orig_idx < len(text):
            if text[orig_idx] == " " or text[orig_idx] == "\n" or text[orig_idx] == "\t":
                while orig_idx < len(text) and text[orig_idx] in " \n\t\r":
                    orig_idx += 1
                norm_idx += 1
            else:
                orig_idx += 1
                norm_idx += 1
        # Extract from orig_idx forward to cover len(norm_excerpt) tokens
        end_idx = orig_idx
        covered = 0
        while covered < len(norm_excerpt) and end_idx < len(text):
            if text[end_idx] in " \n\t\r":
                while end_idx < len(text) and text[end_idx] in " \n\t\r":
                    end_idx += 1
                covered += 1
            else:
                end_idx += 1
                covered += 1
        original_span = text[orig_idx:end_idx]
        if original_span and original_span in text:
            return text.replace(original_span, replacement, 1)
        # Fallback: do the replacement on the normalised string
        replaced_norm = norm_text.replace(norm_excerpt, replacement, 1)
        # Re-normalise the original text and return the replaced version
        return re.sub(r"\s+", " ", replaced_norm)

    # 3. Sentence-level fuzzy match
    sentences = re.split(r"(?<=[.!?\n])\s*", text)
    best_match = None
    best_ratio = 0.0
    for sentence in sentences:
        if not sentence.strip():
            continue
        ratio = SequenceMatcher(None, normalize(sentence), norm_excerpt).ratio()
        if ratio > best_ratio:
            best_ratio = ratio
            best_match = sentence

    if best_match and best_ratio >= 0.85:
        return text.replace(best_match, replacement, 1)

    # 4. Word-token fuzzy match (wider window — whole text chunks)
    window_size = max(1, len(norm_excerpt.split()))
    words = text.split()
    best_window_ratio = 0.0
    best_window_start = -1
    for i in range(len(words) - window_size + 1):
        window = " ".join(words[i : i + window_size])
        ratio = SequenceMatcher(None, normalize(window), norm_excerpt).ratio()
        if ratio > best_window_ratio:
            best_window_ratio = ratio
            best_window_start = i

    if best_window_ratio >= 0.80 and best_window_start >= 0:
        original_window = " ".join(words[best_window_start : best_window_start + window_size])
        if original_window in text:
            return text.replace(original_window, replacement, 1)

    return text

def _replace_in_docx(data: bytes, replacements: List[Tuple[str, str]]) -> bytes:
    import docx

    document = docx.Document(io.BytesIO(data))

    def _apply_to_paragraph(paragraph):
        runs = paragraph.runs
        if not runs:
            return
        full_text = "".join(run.text for run in runs)
        if not full_text:
            return

        # Run boundaries as (start, end) char offsets into full_text, so a
        # matched span can be mapped back to the specific run(s) it lives
        # in — everything outside that span, in this or any other run/
        # paragraph, stays byte-for-byte untouched (text, bold/italic,
        # style, numbering, headings, etc.).
        bounds = []
        pos = 0
        for run in runs:
            start = pos
            pos += len(run.text)
            bounds.append((start, pos))

        for excerpt, replacement in replacements:
            if not excerpt:
                continue
            span = find_source_span(full_text, excerpt)
            if not span:
                continue
            start = full_text.find(span)
            if start == -1:
                continue
            end = start + len(span)

            first_idx = next((i for i, (s, e) in enumerate(bounds) if s <= start < e), None)
            last_idx = next((i for i, (s, e) in enumerate(bounds) if s < end <= e), None)
            if first_idx is None or last_idx is None:
                continue

            if first_idx == last_idx:
                # Span fits inside a single run: edit only that run's text,
                # preserving its formatting and every other run untouched.
                run = runs[first_idx]
                s, e = bounds[first_idx]
                local_start, local_end = start - s, end - s
                run.text = run.text[:local_start] + replacement + run.text[local_end:]
            else:
                # Span crosses a run boundary: the replacement can only take
                # on one run's formatting, so it's applied to the first
                # affected run (keeping that run's own style) and the
                # spanned-over portion of subsequent runs is trimmed —
                # runs before/after the span, and all other paragraphs,
                # are never touched.
                s0, e0 = bounds[first_idx]
                runs[first_idx].text = runs[first_idx].text[: start - s0] + replacement
                for i in range(first_idx + 1, last_idx + 1):
                    si, ei = bounds[i]
                    keep_from = max(0, end - si)
                    runs[i].text = runs[i].text[keep_from:]

            # Recompute full_text/bounds so subsequent replacements in this
            # same paragraph match against the just-updated content.
            full_text = "".join(r.text for r in runs)
            pos = 0
            bounds = []
            for run in runs:
                start_b = pos
                pos += len(run.text)
                bounds.append((start_b, pos))

    def _walk_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _apply_to_paragraph(p)
                    _walk_tables(cell.tables)

    for p in document.paragraphs:
        _apply_to_paragraph(p)
    _walk_tables(document.tables)

    for section in document.sections:
        for p in section.header.paragraphs:
            _apply_to_paragraph(p)
        for p in section.footer.paragraphs:
            _apply_to_paragraph(p)

    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def _replace_in_pdf(data: bytes, replacements: List[Tuple[str, str]]) -> bytes:
    import fitz  # PyMuPDF

    doc = fitz.open(stream=data, filetype="pdf")
    for page in doc:
        for excerpt, replacement in replacements:
            excerpt = (excerpt or "").strip()
            if not excerpt:
                continue
            areas = page.search_for(excerpt)
            for rect in areas:
                page.add_redact_annot(rect, fill=(1, 1, 1))
            if areas:
                page.apply_redactions()
                # Re-insert the compliant text at the same spot, sized to fit.
                rect = areas[0]
                fontsize = max(6, min(11, rect.height * 0.72))
                page.insert_textbox(
                    rect, replacement, fontsize=fontsize, color=(0, 0, 0)
                )
    out = doc.tobytes()
    doc.close()
    return out


def _replace_in_text(data: bytes, replacements: List[Tuple[str, str]]) -> bytes:
    text = data.decode("utf-8", errors="ignore")
    for excerpt, replacement in replacements:
        if excerpt:
            text = smart_replace(text, excerpt, replacement)
    return text.encode("utf-8")


def _appendix_entries(resolved_violations) -> List[dict]:
    """Normalizes ResolvedViolation objects/dicts into the appendix fields.

    Deliberately excludes the original (pre-remediation) clause text: the
    downloaded document must contain only resolved, compliant content —
    never a copy of the vulnerable text it just fixed. Re-printing the
    original clause here would also cause the scanner to re-detect it as a
    fresh violation on a future re-upload of this same file.
    """
    entries = []
    for rv in resolved_violations or []:
        get = (lambda k, d=None: rv.get(k, d)) if isinstance(rv, dict) else (lambda k, d=None: getattr(rv, k, d))
        entries.append({
            "rewrite": get("remediated_text", "") or "",
            "regulation": get("violated_rule", "") or "",
            "reason": get("recommendation", "") or "",
            "status": (get("resolution_status", "resolved") or "resolved").capitalize(),
        })
    return entries


def _append_appendix_docx(data: bytes, resolved_violations) -> bytes:
    import docx

    document = docx.Document(io.BytesIO(data))
    document.add_page_break()
    document.add_heading("Remediation Appendix", level=1)

    for i, e in enumerate(_appendix_entries(resolved_violations), start=1):
        document.add_heading(f"Item {i}", level=2)
        document.add_paragraph(f"Resolved Clause: {e['rewrite']}")
        document.add_paragraph("↓")
        document.add_paragraph(f"Violated Regulation: {e['regulation']}")
        document.add_paragraph("↓")
        document.add_paragraph(f"Reason: {e['reason']}")
        document.add_paragraph("↓")
        document.add_paragraph(f"Resolution Status: {e['status']}")

    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def _append_appendix_pdf(data: bytes, resolved_violations) -> bytes:
    import fitz

    doc = fitz.open(stream=data, filetype="pdf")
    page = doc.new_page()
    y = 50
    lines = ["Remediation Appendix", ""]
    for i, e in enumerate(_appendix_entries(resolved_violations), start=1):
        lines += [
            f"Item {i}",
            f"Resolved Clause: {e['rewrite']}",
            "↓",
            f"Violated Regulation: {e['regulation']}",
            "↓",
            f"Reason: {e['reason']}",
            "↓",
            f"Resolution Status: {e['status']}",
            "",
        ]
    for line in lines:
        if y > page.rect.height - 50:
            page = doc.new_page()
            y = 50
        page.insert_text((50, y), line, fontsize=10)
        y += 14

    out = doc.tobytes()
    doc.close()
    return out


def _append_appendix_text(data: bytes, resolved_violations) -> bytes:
    text = data.decode("utf-8", errors="ignore")
    parts = [text, "\n\n---\nRemediation Appendix\n"]
    for i, e in enumerate(_appendix_entries(resolved_violations), start=1):
        parts.append(
            f"\nItem {i}\n"
            f"Resolved Clause: {e['rewrite']}\n"
            f"↓\n"
            f"Violated Regulation: {e['regulation']}\n"
            f"↓\n"
            f"Reason: {e['reason']}\n"
            f"↓\n"
            f"Resolution Status: {e['status']}\n"
        )
    return "".join(parts).encode("utf-8")


def append_remediation_appendix(data: bytes, ext: str, resolved_violations) -> bytes:
    """
    Takes the already-generated (latest) remediated document's bytes and
    appends a Remediation Appendix — one entry per resolved violation, in
    Original Clause -> AI Rewrite -> Violated Regulation -> Reason ->
    Resolution Status order. Used only for the download response; the
    document a user *previews* is left untouched by this.
    """
    ext = (ext or "").lower().lstrip(".")
    try:
        if ext == "docx":
            return _append_appendix_docx(data, resolved_violations)
        if ext == "pdf":
            return _append_appendix_pdf(data, resolved_violations)
        if ext in ("txt", "md"):
            return _append_appendix_text(data, resolved_violations)
    except Exception as e:
        print(f"[doc_remediator] Failed to append remediation appendix: {e}")
    return data


def merge_overlapping_replacements(replacements: List[Tuple[str, str]]) -> List[Tuple[str, str]]:
    """
    Multiple violations can share the exact same excerpt (e.g. one sentence
    that is simultaneously an unsafe cross-border-transfer clause AND an
    indefinite-retention clause). Treating each as an independent full-line
    replacement causes the second replacement to run against text the first
    one already changed — the exact excerpt is gone, so it falls through to
    fuzzy matching and can revert or duplicate the first fix.

    Instead, replacements that target the same original excerpt are merged:
    each one's specific changed fragment (computed via a diff against the
    shared excerpt) is layered onto a single working copy, in order, so both
    fixes land in the final text without clobbering each other.
    """
    working_by_excerpt: dict = {}
    order: List[str] = []

    for excerpt, replacement in replacements:
        key = (excerpt or "").strip()
        replacement = (replacement or "").strip()
        if not key or not replacement:
            continue
        if key not in working_by_excerpt:
            working_by_excerpt[key] = key
            order.append(key)

        if replacement == key:
            continue

        working = working_by_excerpt[key]
        if working == key:
            # First fix for this excerpt: take the remediated text as-is.
            working_by_excerpt[key] = replacement
            continue

        # Subsequent fix for an excerpt already modified by a prior
        # violation: apply only this violation's changed fragment(s) on
        # top of the current working copy, so earlier fixes are preserved.
        # Diffed on whitespace-preserving word tokens (not raw characters)
        # so fragment boundaries always land on whole words.
        key_tokens = re.findall(r"\S+|\s+", key)
        repl_tokens = re.findall(r"\S+|\s+", replacement)
        sm = SequenceMatcher(None, key_tokens, repl_tokens)
        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            if tag == "equal":
                continue
            old_frag = "".join(key_tokens[i1:i2])
            new_frag = "".join(repl_tokens[j1:j2])
            if old_frag and old_frag in working:
                working = working.replace(old_frag, new_frag, 1)
        working_by_excerpt[key] = working

    return [(k, working_by_excerpt[k]) for k in order if working_by_excerpt[k] != k]


def regenerate_document(original_path: str, ext: str, replacements: List[Tuple[str, str]]) -> bytes:
    """
    Reads the real uploaded file from disk and returns new bytes with the
    given (excerpt -> remediated_text) replacements applied in place.
    Falls back to returning the untouched original bytes if the format
    isn't editable or nothing could be matched.
    """
    data = Path(original_path).read_bytes()
    ext = (ext or "").lower().lstrip(".")
    replacements = merge_overlapping_replacements(replacements)

    try:
        if ext == "docx":
            return _replace_in_docx(data, replacements)
        if ext == "pdf":
            return _replace_in_pdf(data, replacements)
        if ext in ("txt", "md"):
            return _replace_in_text(data, replacements)
    except Exception as e:
        print(f"[doc_remediator] Failed to regenerate {ext} document: {e}")

    return data
